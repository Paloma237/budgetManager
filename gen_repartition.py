# -*- coding: utf-8 -*-
"""Genere le document Word de repartition des taches pour BudgetManager."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----- Couleurs (alignees sur le CDC) -----
BLEU_FONCE = RGBColor(0x1F, 0x38, 0x64)
BLEU = RGBColor(0x2E, 0x75, 0xB6)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)
GRIS = RGBColor(0x59, 0x59, 0x59)

HEADER_FILL = "2E75B6"
ALT_FILL = "DEEAF6"

doc = Document()

# Marges
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

# Police par defaut
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, align="left", italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color is not None:
        run.font.color.rgb = color


def add_table(headers, rows, widths=None, header_fill=HEADER_FILL, alt=True, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # entete
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], header_fill)
        set_cell_text(hdr[i], h, bold=True, color=BLANC, size=font_size, align="center")
    # lignes
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if alt and r_idx % 2 == 1:
                set_cell_bg(cells[i], ALT_FILL)
            align = "center" if (len(str(val)) <= 4) else "left"
            set_cell_text(cells[i], str(val), size=font_size, align=align)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def h1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BLEU_FONCE
    # bordure basse
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2E75B6")
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    return p


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = BLEU
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def para(text, italic=False, size=11, color=None, align="justify"):
    p = doc.add_paragraph()
    p.alignment = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


# =====================================================================
# PAGE DE GARDE
# =====================================================================
def center_line(text, size, color, bold=True, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

center_line("UNIVERSITÉ DE DSCHANG", 14, BLEU_FONCE)
center_line("IUT Fotso Victor de Bandjoun — Département Génie Informatique", 11, GRIS)
doc.add_paragraph()
doc.add_paragraph()
center_line("DOCUMENT DE GESTION DE PROJET", 20, BLEU_FONCE)
center_line("Cycle de développement & Répartition des tâches", 14, BLEU)
doc.add_paragraph()
center_line("Application Mobile Android « BudgetManager »", 16, BLEU_FONCE)
doc.add_paragraph()
doc.add_paragraph()

# tableau identite
info_rows = [
    ["Module", "Développement d'applications pour terminaux mobiles"],
    ["Enseignant", "Dr. JAGHO MDEMAYA G. Brel"],
    ["Établissement", "IUT Fotso Victor de Bandjoun"],
    ["Année académique", "2025 – 2026"],
    ["Effectif du groupe", "4 membres"],
    ["Date", "Juin 2026"],
]
t = add_table(["Champ", "Détail"], info_rows, widths=[5, 11], font_size=11)

doc.add_page_break()

# =====================================================================
# 1. OBJET DU DOCUMENT
# =====================================================================
h1("1. Objet du document")
para("Ce document accompagne le cahier des charges de l'application BudgetManager. "
     "Il a pour but de structurer le projet selon les grandes étapes du cycle de "
     "développement d'un logiciel — de l'analyse des besoins jusqu'à la documentation — "
     "et de répartir clairement les tâches entre les trois membres du groupe.")
para("Conformément aux consignes, l'étape de déploiement (publication/distribution) "
     "n'est pas couverte : le projet s'arrête à une application fonctionnelle, testée "
     "sur émulateur AVD et documentée. Le périmètre respecte strictement les exigences "
     "du cahier des charges ; les améliorations éventuelles (graphiques, export, etc.) "
     "sont reportées à une version ultérieure.", italic=False)

# =====================================================================
# 2. EQUIPE PROJET
# =====================================================================
h1("2. Présentation de l'équipe projet")
para("Le groupe est composé de trois membres. Chaque membre pilote un ou plusieurs axes "
     "du projet, mais l'ensemble de l'équipe participe à la relecture du code et à la "
     "préparation de la soutenance.")

equipe_rows = [
    ["Kamogne Tedonnang", "KT", "Chef de projet / Analyse (A) · Logique applicative (E)",
     "Coordination, analyse des besoins, planning, codage de MainActivity.java"],
    ["MOTSO SIMO PALOMA BATISTA", "MP", "Conception & modélisation (B) · Base de données (C)",
     "Diagrammes UML, modèle de données, codage de DatabaseHelper.java"],
    ["KANA KENMENE Laura Princesse", "LP", "Maquettes · Interface (D)",
     "Maquettes de l'écran, activity_main.xml / strings.xml / colors.xml"],
    ["MOTOUOM MBA ISABELLE SANDRA", "IS", "Tests & Documentation (F)",
     "Plan de tests, exécution sur AVD, rapport de projet, manuel d'utilisation"],
]
add_table(["Membre", "Code", "Rôles", "Responsabilités principales"],
          equipe_rows, widths=[4.2, 1.3, 4.5, 6], font_size=9.5)

para(" ")
para("Les codes KT, MP, LP et IS sont utilisés dans les tableaux de répartition et la "
     "matrice RACI ci-après.", italic=True, size=10, color=GRIS)

# =====================================================================
# 3. CYCLE DE DEVELOPPEMENT ADOPTE
# =====================================================================
h1("3. Cycle de développement adopté")
para("Le groupe adopte un modèle de développement séquentiel (modèle en cascade), "
     "bien adapté à un projet de taille réduite, au périmètre figé par le cahier des "
     "charges et au délai court. Les étapes s'enchaînent logiquement, chacune produisant "
     "un livrable qui alimente la suivante.")

cycle_rows = [
    ["1", "Analyse des besoins", "Comprendre et formaliser ce que l'application doit faire (exigences F1–F5)."],
    ["2", "Conception (spécifications)", "Modéliser la solution : UML, modèle de données, maquettes de l'interface."],
    ["3", "Conception détaillée", "Définir l'architecture (2 classes) et les signatures des méthodes."],
    ["4", "Implémentation (codage)", "Écrire le code : interface XML, base SQLite, logique applicative."],
    ["5", "Tests & validation", "Vérifier chaque fonctionnalité sur l'émulateur AVD."],
    ["6", "Documentation & soutenance", "Rédiger le rapport, le manuel utilisateur et préparer la démonstration."],
]
add_table(["#", "Étape", "Description"], cycle_rows, widths=[1, 4.5, 10.5], font_size=10)
para(" ")
para("Note : l'étape de déploiement (mise en production / distribution de l'APK au "
     "public) est volontairement exclue du périmètre de ce projet.", italic=True, size=10, color=GRIS)

# =====================================================================
# 4. ETAPES DETAILLEES + TACHES
# =====================================================================
h1("4. Détail des étapes et des tâches")

def etape(titre, intro, tasks):
    h2(titre)
    if intro:
        para(intro)
    add_table(["Tâche", "Description", "Responsable"], tasks,
              widths=[4.5, 9, 2.5], font_size=9.5)
    para(" ")

etape("Étape 1 — Analyse des besoins",
      "Objectif : transformer le cahier des charges en exigences claires et vérifiables.",
      [["Reformulation du CDC", "Relire le cahier des charges et lister les exigences fonctionnelles (F1 à F5).", "KT"],
       ["Exigences non fonctionnelles", "Identifier les contraintes : Java, SQLite, API 16, hors-ligne, aucune permission.", "KT"],
       ["Glossaire & règles de gestion", "Définir : Solde = Total revenus − Total dépenses ; types « depense »/« revenu ».", "KT"],
       ["Validation des besoins", "Valider la liste des besoins avec toute l'équipe.", "KT, MP, LP"]])

etape("Étape 2 — Conception (spécifications)",
      "Objectif : modéliser la solution avant de coder.",
      [["Diagramme de cas d'utilisation", "Acteur « Utilisateur » et cas : ajouter, consulter solde, lister, supprimer.", "MP"],
       ["Diagramme de classes", "Modéliser MainActivity et DatabaseHelper et leurs relations.", "MP"],
       ["Modèle de données", "Concevoir la table transactions (id, montant, type, description, date).", "MP"],
       ["Maquettes de l'interface", "Dessiner l'écran unique : zone solde, saisie, bouton, historique.", "LP"]])

etape("Étape 3 — Conception détaillée",
      "Objectif : préparer le terrain technique du codage.",
      [["Architecture applicative", "Confirmer l'architecture à 2 classes (MainActivity + DatabaseHelper).", "KT"],
       ["Signatures BDD", "Définir ajouterTransaction(), getToutesTransactions(), supprimerTransaction().", "MP"],
       ["Signatures logique", "Définir calculerSolde(), actualiserListe() et les listeners.", "KT"],
       ["Conventions de code", "Fixer le nom du package, les conventions de nommage et les IDs de widgets.", "KT, MP, LP"]])

etape("Étape 4 — Implémentation (codage)",
      "Objectif : produire le code conforme au cahier des charges.",
      [["Interface — activity_main.xml", "LinearLayout vertical avec tous les widgets (IDs imposés par le CDC).", "LP"],
       ["Ressources — strings.xml / colors.xml", "Centraliser les libellés et la palette de couleurs.", "LP"],
       ["Base de données — DatabaseHelper.java", "SQLiteOpenHelper : onCreate, onUpgrade, CRUD via ContentValues/Cursor.", "MP"],
       ["Logique — MainActivity.java", "Listeners, calcul du solde, affichage ListView, suppression par appui long.", "KT"],
       ["Intégration", "Assembler UI + BDD + logique et résoudre les anomalies d'intégration.", "KT, MP"]])

etape("Étape 5 — Tests & validation",
      "Objectif : garantir que chaque fonctionnalité marche sur l'émulateur AVD.",
      [["Plan de tests", "Rédiger les cas de test pour F1 à F5 (résultats attendus).", "IS"],
       ["Jeu de données de test", "Préparer des transactions types (revenus et dépenses).", "IS"],
       ["Exécution des tests", "Tester ajout, solde, historique, suppression et persistance sur AVD.", "IS, KT"],
       ["Correction des anomalies", "Corriger les bugs détectés (BDD ou logique).", "MP, KT"]])

etape("Étape 6 — Documentation & soutenance",
      "Objectif : livrer un projet documenté et prêt à présenter.",
      [["Rapport de projet", "Rédiger le rapport (démarche, captures, choix techniques).", "IS"],
       ["Manuel d'utilisation", "Expliquer comment utiliser l'application pas à pas.", "IS"],
       ["Génération de l'APK de démo", "Produire l'APK installable pour la démonstration.", "KT"],
       ["Préparation de la soutenance", "Préparer le support et répéter la démonstration en équipe.", "KT, MP, LP, IS"]])

# =====================================================================
# 5. SYNTHESE PAR MEMBRE
# =====================================================================
doc.add_page_break()
h1("5. Synthèse de la répartition par membre")

synth_rows = [
    ["Kamogne Tedonnang (KT)",
     "• Coordination du projet et suivi du planning\n"
     "• Analyse des besoins et règles de gestion\n"
     "• Conception détaillée de la logique\n"
     "• Codage de MainActivity.java (listeners, calcul du solde, ListView)\n"
     "• Génération de l'APK de démonstration"],
    ["MOTSO SIMO PALOMA BATISTA (MP)",
     "• Diagrammes UML (cas d'utilisation, classes)\n"
     "• Conception du modèle de données (table transactions)\n"
     "• Codage de DatabaseHelper.java (SQLiteOpenHelper + CRUD)\n"
     "• Correction des anomalies liées à la base de données"],
    ["KANA KENMENE Laura Princesse (LP)",
     "• Maquettes de l'interface\n"
     "• Codage de activity_main.xml, strings.xml, colors.xml\n"
     "• Ergonomie et respect des IDs de widgets imposés par le CDC"],
    ["MOTOUOM MBA ISABELLE SANDRA (IS)",
     "• Plan de tests des fonctionnalités F1 à F5\n"
     "• Exécution des tests sur émulateur AVD\n"
     "• Rédaction du rapport de projet\n"
     "• Rédaction du manuel d'utilisation"],
]
t = doc.add_table(rows=0, cols=2)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.add_row().cells
for i, htxt in enumerate(["Membre", "Contributions"]):
    set_cell_bg(hdr[i], HEADER_FILL)
    set_cell_text(hdr[i], htxt, bold=True, color=BLANC, size=11, align="center")
for r_idx, (m, c) in enumerate(synth_rows):
    cells = t.add_row().cells
    if r_idx % 2 == 1:
        set_cell_bg(cells[0], ALT_FILL)
        set_cell_bg(cells[1], ALT_FILL)
    set_cell_text(cells[0], m, bold=True, size=10, align="left")
    set_cell_text(cells[1], c, size=10, align="left")
    cells[0].width = Cm(5)
    cells[1].width = Cm(11)

# =====================================================================
# 6. MATRICE RACI
# =====================================================================
h1("6. Matrice de responsabilités (RACI)")
para("Légende — R : Réalise (exécute la tâche) · A : Approuve / responsable final · "
     "C : Consulté · I : Informé.", italic=True, size=10, color=GRIS)

raci_rows = [
    ["Analyse des besoins", "A,R", "C", "C", "C"],
    ["Modélisation UML", "A", "R", "I", "I"],
    ["Modèle de données (table transactions)", "A", "R", "I", "I"],
    ["Maquettes de l'interface", "C", "C", "R", "I"],
    ["Conception détaillée", "A,R", "C", "I", "I"],
    ["activity_main.xml / strings / colors", "I", "C", "R", "C"],
    ["DatabaseHelper.java (SQLite)", "C", "R", "I", "I"],
    ["MainActivity.java (logique)", "R", "C", "I", "I"],
    ["Intégration", "A,R", "R", "C", "C"],
    ["Tests sur émulateur AVD", "C", "C", "C", "R"],
    ["Correction des anomalies", "R", "R", "I", "C"],
    ["Rapport & manuel d'utilisation", "C", "C", "C", "R"],
    ["Génération de l'APK de démo", "R", "I", "I", "I"],
    ["Préparation de la soutenance", "A,R", "R", "R", "R"],
]
add_table(["Tâche / Livrable", "KT", "MP", "LP", "IS"], raci_rows,
          widths=[8.2, 2, 2, 2, 2], font_size=9.5)

# =====================================================================
# 7. PLANNING
# =====================================================================
h1("7. Planning de réalisation")
para("Compte tenu du délai court (soutenance dans la semaine), le planning est resserré "
     "sur sept jours. Les étapes peuvent légèrement se chevaucher (par exemple, les "
     "maquettes commencent dès la fin de l'analyse).")

planning_rows = [
    ["J1", "Analyse des besoins + lancement de la conception", "KT (MP, LP, IS)"],
    ["J2", "Modélisation UML, modèle de données, maquettes", "MP, LP"],
    ["J3", "Conception détaillée + démarrage du codage + plan de tests", "KT, MP, IS"],
    ["J4", "Codage : interface (LP) et base de données (MP)", "LP, MP"],
    ["J5", "Codage : logique MainActivity + intégration", "KT (MP)"],
    ["J6", "Tests sur AVD et correction des anomalies", "IS, KT, MP"],
    ["J7", "Documentation, APK de démo, préparation soutenance", "IS, KT, MP, LP"],
]
add_table(["Jour", "Activité principale", "Membres mobilisés"], planning_rows,
          widths=[1.6, 9.4, 5], font_size=10)

# =====================================================================
# 8. LIVRABLES
# =====================================================================
h1("8. Livrables attendus (hors déploiement)")
livr_rows = [
    ["L1", "Cahier des charges", "Document de référence (déjà fourni).", "—"],
    ["L2", "Dossier de conception", "Diagrammes UML, modèle de données, maquettes.", "MP, LP"],
    ["L3", "Code source complet", "MainActivity.java, DatabaseHelper.java, activity_main.xml, ressources.", "KT, MP, LP"],
    ["L4", "Application fonctionnelle", "APK installable + démonstration sur émulateur AVD.", "KT"],
    ["L5", "Plan de tests", "Cas de test des fonctionnalités F1 à F5 et résultats.", "IS"],
    ["L6", "Rapport & manuel", "Rapport de projet et manuel d'utilisation.", "IS"],
]
add_table(["#", "Livrable", "Description", "Responsable"], livr_rows,
          widths=[1, 4, 8.5, 2.5], font_size=9.5)

para(" ")
para(" ")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Fait à Bandjoun, Juin 2026")
run.italic = True
run.font.size = Pt(11)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r2 = p2.add_run("Le groupe projet : Kamogne Tedonnang · MOTSO SIMO Paloma Batista · "
                "KANA KENMENE Laura Princesse · MOTOUOM MBA Isabelle Sandra")
r2.italic = True
r2.font.size = Pt(10)
r2.font.color.rgb = GRIS

out = r"C:\Users\Arnol\Desktop\BudgetManager\BudgetManager_Repartition_Taches.docx"
doc.save(out)
print("Document genere :", out)
